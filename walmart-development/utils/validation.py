from io import BytesIO

import requests
from docx import Document


def load_sop_from_url(url_or_path: str) -> str:
    if url_or_path.startswith(("http://", "https://")):
        response = requests.get(url_or_path)
        response.raise_for_status()
        doc = Document(BytesIO(response.content))
    else:
        doc = Document(url_or_path)

    parts = []
    parts.append("=== AE VALIDATION SOP ===")
    parts.append("")

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else "Normal"
        if "Heading" in style or "Title" in style:
            level = style.replace("Heading ", "").replace("Title", "1")
            prefix = "#" * int(level) if level.isdigit() else "##"
            parts.append(f"{prefix} {text}")
            parts.append("")
        elif "List" in style:
            parts.append(f"  - {text}")
        else:
            parts.append(text)
            parts.append("")

    parts.append("")
    parts.append("=== TABLES ===")
    parts.append("")

    for i, table in enumerate(doc.tables, 1):
        parts.append(f"--- Table {i} ---")
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells)
        if rows_data:
            header = " | ".join(rows_data[0])
            parts.append(f"[Header] {header}")
            for row in rows_data[1:]:
                parts.append(f"  { ' | '.join(row) }")
        parts.append("")

    return "\n".join(parts)
